"""
Funções da Interface Grafica
"""
from PIL import Image, ImageTk
from datetime import datetime
from docx import Document
from python_docx_replace import docx_replace
from ttkbootstrap.constants import *
from ManagerBD import BancoDados


class Funcs(object):

    def ativa_caucao(self):
        if self.caucaoVar.get() == 1:
            self.txtCaucao.grid(row=7, column=1, sticky='W', padx=10)
            self.caucao.grid(row=8, column=1, sticky='EW', padx=10, columnspan=2)
        elif self.caucaoVar.get() == 0:
            self.txtCaucao.grid_forget()
            self.caucao.grid_forget()


    def add_fiador(self):

        if self.varfiad.get() == 1:
            self.notbook.add(self.fiador)
        
        else:
            self.notbook.hide(2)

    # Deixar a Entry com o formato padrão de RG: xx.xxx.xxx-xx
    def format_rg(self, event):
    
        self.rgtext = event.widget.get().replace(".", "").replace("-", "")[:9]
        self.rgnew_text = ""

        if event.keysym.lower() == "backspace": 
            return
        
        for index in range(len(self.rgtext)):
            
            if not self.rgtext[index] in "0123456789": 
                continue
            if index in [1, 4]: 
                self.rgnew_text += self.rgtext[index] + "."
            elif index == 7: 
                self.rgnew_text += self.rgtext[index] + "-"
            else: 
                self.rgnew_text += self.rgtext[index]

        event.widget.delete(0, "end")
        event.widget.insert(0, self.rgnew_text)

    # Deixar a Entry com o formato padrão de CPF: xxx.xxx.xxx-xx
    def format_cpf(self, event):
    
        self.text = event.widget.get().replace(".", "").replace("-", "")[:11]
        self.new_text = ""

        if event.keysym.lower() == "backspace": 
            return
        
        for index in range(len(self.text)):
            
            if not self.text[index] in "0123456789": 
                continue
            if index in [2, 5]: 
                self.new_text += self.text[index] + "."
            elif index == 8: 
                self.new_text += self.text[index] + "-"
            else: 
                self.new_text += self.text[index]

        event.widget.delete(0, "end")
        event.widget.insert(0, self.new_text)

    # Para habilitar e desabilitar o frame do segundo locatario quando ouver
    def ativa_locatarios(self, ql_frame):
        if ql_frame == 1:

            if self.quantidadeLocatarios.get() == 1:

                for child in self.frameMora2.winfo_children():
                    child.configure(state='disable')

            elif self.quantidadeLocatarios.get() == 2:

                for child in self.frameMora2.winfo_children():
                    child.configure(state='normal')
                    self.combMora2.config(state='readonly')

        elif ql_frame == 2:

            if self.var2fiador.get() == 0:

                for child in self.lb_fram_f2.winfo_children():
                    child.configure(state='disable')

            elif self.var2fiador.get() == 1:

                for child in self.lb_fram_f2.winfo_children():
                    child.configure(state='normal')
                    self.combFiador2.config(state='readonly')    

    def limp_tela(self, ttk):
        """
        OPÇÃO ALTERNATIVA PARA LIMPAR AS ENTRY
         [widget.delete(0, END) for widget in self.frameContrato.winfo_children() if isinstance(widget, ttk.Entry)]
         [widget.delete(0, END) for widget in self.frameMora1.winfo_children() if isinstance(widget, ttk.Entry)]
         [widget.delete(0, END) for widget in self.frameMora2.winfo_children() if isinstance(widget, ttk.Entry)]
         [widget.delete(0, END) for widget in self.frameLocador.winfo_children() if isinstance(widget, ttk.Entry)]
         [widget.delete(0, END) for widget in self.fiador.winfo_children() if isinstance(widget, ttk.Entry)]
         [widget.delete(0, END) for widget in self.fiador2.winfo_children() if isinstance(widget, ttk.Entry)]

        """
        conf_msg = ttk.dialogs.dialogs.Messagebox.yesno("Confirme para limpar todas as caixas de entrada?", "Confimação")
        if conf_msg == 'Yes':
            frames = [self.lb_fram_f1, self.lb_fram_f2, self.lb_fram_end, self.frameLocador, self.frameMora1, self.frameMora2, self.frameContrato]
            for fm in frames:
                for widget in self.root.winfo_children():
                    if not isinstance(widget, ttk.Entry):
                        #clear(widget)
                        [widget.delete(0, END) for widget in fm.winfo_children() if isinstance(widget, ttk.Entry)]
                    elif isinstance(widget, ttk.Entry):
                        widget.delete(0, ttk.END)

    def variaveis_images(self):

        self.f_on_image = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/feminino.png").resize((78,78), resample=3)
        self.f_on_image = ImageTk.PhotoImage(self.f_on_image)
        
        self.h_on_image = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/sexo-masculino.png").resize((78,78), resample=3)
        self.h_on_image = ImageTk.PhotoImage(self.h_on_image)
        
        self.imgCasa = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/casa.png").resize((20,20), resample=3)
        self.imgCasa = ImageTk.PhotoImage(self.imgCasa)

        img_add = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/pesquisa-de-pessoas.png").resize((25,25), resample=3)
        self.add_locatarios_img = ImageTk.PhotoImage(img_add)

        img_t = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/LOGO_PRETA.png")
        img_p = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/LOGO_RM.png")
        self.img = img_t.resize((250,250), resample=3)
        self.fundo = ImageTk.PhotoImage(self.img)
        self.img_peq = img_p.resize((130,55), resample=3)
        self.fund_pq = ImageTk.PhotoImage(self.img_peq)

 
        self.addImg = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/add.png").resize((25,25), resample=3)
        self.addImg = ImageTk.PhotoImage(self.addImg)

        self.addImg2 = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/add.png").resize((23,23), resample=3)
        self.addImg2 = ImageTk.PhotoImage(self.addImg2)

        lix = Image.open("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/imgs/lixeira.png").resize((20,20), resample=3)
        self.lixeira = ImageTk.PhotoImage(lix)

    def alte_img(self, event):
        wid_press = str(event.widget.info).replace('.', '').split('!')

        if wid_press[2] == 'frame':

            if self.combStr.get() == 'Homem':
                self.img_locador['image'] = self.h_on_image
                self.nameLa.config(bootstyle=PRIMARY)

            else:
                self.img_locador['image'] = self.f_on_image
                self.nameLa.config(bootstyle=DANGER)

        elif wid_press[2] == 'frame2' and wid_press[3] == 'frame':

            if self.combMoraStr.get() == 'Homem':
                self.img_mora['image'] = self.h_on_image

            else:
                self.img_mora['image'] = self.f_on_image
        
        elif wid_press[2] == 'frame2' and wid_press[3] == 'frame2':

            if self.combMora2Str.get() == 'Homem':
                self.img_mora2['image'] = self.h_on_image

            else:
                self.img_mora2['image'] = self.f_on_image

        elif wid_press[2] == 'frame3':
            
            if wid_press[3] == 'labelframe':

                if self.combFiadorStr.get() == 'Homem':
                    self.img_fiador['image'] = self.h_on_image

                else:
                    self.img_fiador['image'] = self.f_on_image
            
            else:
                if self.combFiadorStr2.get() == 'Homem':
                    self.img_fiador2['image'] = self.h_on_image

                else:
                    self.img_fiador2['image'] = self.f_on_image

    def salvar_word(self):
        # mulher 1 Homem 0
        date_obj = datetime.strptime(self.inicioImov.entry.get(), '%d/%m/%Y')
        date_mes = datetime.strftime(date_obj, "%B")

        if self.iptuVar.get() == 0:
            pagamento_iptu = ('LOCATÁRIA', 'a') if self.combMoraStr.get() == 'Mulher' else ('LOCATÁRIO', 'o')
        else:
            pagamento_iptu = ('LOCADORA', 'a') if self.combStr.get() == 'Mulher' else ('LOCADOR', 'o')

        data_list = self.inicioImov.entry.get().split("/")
        loca_dict = {
            "pron1": 'A' if self.combStr.get() == 'Mulher' else '',
            "pron1.11": "a" if self.combStr.get() == 'Mulher' else 'o',
            "pron1.2": "A" if self.combStr.get() == 'Mulher' else 'O',
            "nome1": str(self.nomeLoca.get()),
            "est civil1": str(self.estcivLoca.get()),
            "pron1.1": 'a' if self.combStr.get() == 'Mulher' else '',
            "rg1": str(self.rgLoca.get()),
            "org espeditor1": str(self.org_ex_locador.get()),
            "estado1": str(self.estado_rg_locador.get()),
            "cpf1": str(self.cpfLoca.get()),
            "municipio1": str(self.cidad_locador.get()),
            "estado1.1": str(self.est_locador.get()),
            "rua ou avenida1": str(self.ruaLocador.get()),
            "numero1": str(self.numeroLocador.get()),
            "complemento1": str(self.complLocador.get()+', ') if self.complLocador.get() else '',
            "bairro1": str(self.bairroLocador.get()),
            "cep1": str(self.cepLocador.get())
        }
        morador_dict = {
            "pron2": 'A' if self.combMoraStr.get() == 'Mulher' else 'O',
            "pron2.1": 'a' if self.combMoraStr.get() == 'Mulher' else '',
            "pron2.11": 'a' if self.combMoraStr.get() == 'Mulher' else 'o',
            "nome2": str(self.nomeMora1.get()),
            "est civil2": str(self.estcivMora.get()),
            "rg2": str(self.rgMora.get()),
            "org expeditor2": str(self.org_ex_mora.get()),
            "estado2": str(self.estado_rg_mora.get()),
            "cpf2": str(self.cpfMora.get()),
            "municipio2": str(self.cidad_mora.get()),
            "estado2.1": str(self.est_mora.get()),
            "pron5": 'a' if self.combMoraStr.get() == 'Mulher' else 'ao' 
        }
        morador2_dict = {
            "pron4": 'A' if self.combMora2Str.get() == 'Mulher' else 'O',
            "pron4.1": 'a' if self.combMora2Str.get() == 'Mulher' else 'o',
            "pron4.2": 'a' if self.combMora2Str.get() == 'Mulher' else '',
            "nome4": str(self.nomeMora2.get()),
            "est civil4": str(self.estcivMora2.get()),
            "rg4": str(self.rgMora2.get()),
            "org expeditor4": str(self.org_ex_mora2.get()),
            "estado4": str(self.estado_rg_mora2.get()),
            "cpf4": str(self.cpfMora2.get()),
            "municipio2": str(self.cidad_mora.get()),
            "estado2.1": str(self.est_mora.get()), 
        }
        fiador_dict = {
            "pron3": 'A' if self.combFiadorStr.get() == 'Mulher' else '',
            "pron3.1": 'a' if self.combFiadorStr.get() == 'Mulher' else '',
            "pron3.11": 'a' if self.combFiadorStr.get() == 'Mulher' else 'o',
            "pron3.2": 'A' if self.combFiadorStr.get() == 'Mulher' else 'O',
            "nome3": str(self.nomeFiador.get()),
            "est civil3": str(self.estcivFiador.get()),
            "rg3": str(self.rgFiador.get()),
            "org expeditor3": str(self.org_ex_fiador.get()),
            "estado3": str(self.estado_rg_fiador.get()),
            "cpf3": str(self.cpfFiador.get()),
            "municipio3": str(self.cidad_fiador.get()),
            "estado3.1": str(self.est_fiador.get()),
            "rua ou avenida3": str(self.ruaFiador.get()),
            "numero3": str(self.numeroFiador.get()),
            "complemento3": str(self.complFiador.get()+', ') if self.complFiador.get() else '',
            "bairro3": str(self.bairroFiador.get()),
            "cep3": str(self.cepFiador.get())
        }
        fiador_dict2 = {
            "pron5.1": 'a' if self.combFiadorStr2.get() == 'Mulher' else 'o',
            "pron5.2": 'a' if self.combFiadorStr2.get() == 'Mulher' else '',
            "nome5": str(self.nomeFiador2.get()),
            "est civil5": str(self.estcivFiador2.get()),
            "rg5": str(self.rgFiador2.get()),
            "org expeditor5": str(self.org_ex_fiador2.get()),
            "estado5": str(self.estado_rg_fiador2.get()),
            "cpf5": str(self.cpfFiador2.get()),
        }
        contrato_dic = {
            "rua imov locado": self.endImov.get(),
            "numero4": self.numImov.get(),
            "complemento4": str(self.complImov.get()+', ') if self.complImov.get() else '',
            "bairro4": self.bairImov.get(),
            "municipio4": self.municImov.get(),
            "estado4": self.estImov.get(),
            "finalidade": 'Comercial' if self.finalidadeVar.get() == 1 else 'Residencial',
            "prazo": self.prazo.get(),
            "data inicio": self.inicioImov.entry.get(),
            "data termino": self.fimImov.entry.get()[:10],
            "vencimento": self.vencImov.get(),
            "valor": self.valorImov.get(),
            "valor caução": self.caucao.get(),
            "pg iptu": pagamento_iptu[0],
            "forma pg alug": self.formPG.get(),
            "pron6": pagamento_iptu[1],
            "dia": data_list[0],
            "mês": date_mes,
            "ano": data_list[2]
        }

        self.db = BancoDados()
        print(self.nomeLoca.get())
        self.db.insert_locador([str(self.nomeLoca.get()),
                                str(self.estcivLoca.get()),
                                str(self.cpfLoca.get()),
                                str(self.rgLoca.get()),
                                str(self.org_ex_locador.get()),
                                str(self.estado_rg_locador.get()),
                                str(self.cidad_locador.get()),
                                str(self.est_locador.get()),
                                str(self.numeroLocador.get()),
                                str(self.ruaLocador.get()),
                                str(self.bairroLocador.get())])
        self.db.insert_morador([self.nomeMora1.get(),
                                self.estcivMora.get(),
                                self.cpfMora.get(),
                                self.rgMora.get(),
                                self.org_ex_mora.get(),
                                self.estado_rg_mora.get(),
                                self.cidad_mora.get(),
                                self.est_mora.get()])
        
        tipo_imv = 'Comercial' if self.finalidadeVar.get() == 1 else 'Residencial'
        locador_id_db = self.db.return_ids(str(self.nomeLoca.get()), "Locadores")[0]
        morador_id_db = self.db.return_ids(str(self.nomeMora1.get()), "Locatarios")[0]
        
        self.db.insert_cont([self.endImov.get(),
                             self.numImov.get(),
                             self.municImov.get(),
                             self.estImov.get(),
                             self.bairImov.get(),
                             self.inicioImov.entry.get(),
                             self.fimImov.entry.get()[:10],
                             self.prazo.get(),
                             self.valorImov.get(),
                             self.caucao.get(),
                             pagamento_iptu[0],
                             tipo_imv,
                             locador_id_db, morador_id_db]
                             )
        # Sem Fiador e sem Caução
        if self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 0 and self.varfiad.get() == 0 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | contrato_dic
            tipo = ""
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato01-locatário.docx")
        
        # Sem Fiador e com caução
        elif self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 1 and self.varfiad.get() == 0 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | contrato_dic
            tipo = '(COM CAUÇÃO)'
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato02-locatario-caucao.docx")

        # Com um Locatario, fiador e caução
        elif self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 1 and self.varfiad.get() == 1 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | fiador_dict | contrato_dic
            tipo = "(COM FIADOR) (COM CAUÇÂO)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato03-locatario-fiador-caucao.docx")

        # Com um Locatario e fiador 
        elif self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 0 and self.varfiad.get() == 1 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | fiador_dict | contrato_dic
            tipo = "(COM FIADOR)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato04-locatario-fiador.docx")

        # Com um Locatario e dois fiadores 
        elif self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 0 and self.varfiad.get() == 1 and self.var2fiador.get() == 1:
            my_dict = loca_dict | morador_dict | fiador_dict | fiador_dict2| contrato_dic
            tipo = "(COM FIADOR)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato12-locatario-fiadores.docx")
                                                                                         

        # Com um Locatario e dois fiadores 
        elif self.quantidadeLocatarios.get() == 1 and self.caucaoVar.get() == 1 and self.varfiad.get() == 1 and self.var2fiador.get() == 1:
            my_dict = loca_dict | morador_dict | fiador_dict | fiador_dict2| contrato_dic
            tipo = "(COM FIADOR)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato11-locatario-fiadores-caucao.docx")

        # Dois locatarios
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 0 and self.varfiad.get() == 0 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | fiador_dict | contrato_dic | morador2_dict
            tipo = "TT"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato05-locatarios.docx")

        # Com dois Locatarios e caução
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 1 and self.varfiad.get() == 0:
            my_dict = loca_dict | morador_dict | morador2_dict | contrato_dic
            tipo = "(COM CAUÇÂO)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato06-locatarios-caucao.docx")

        # Com dois Locatarios, um Fiador e caucão
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 1 and self.varfiad.get() == 1 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | morador2_dict | fiador_dict | contrato_dic
            tipo = "(COM FIADOR) (COM CAUÇÂO)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato07-locatarios-fiador-caucao.docx")
        
        # Com dois Locatarios e fiador
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 0 and self.varfiad.get() == 1 and self.var2fiador.get() == 0:
            my_dict = loca_dict | morador_dict | morador2_dict | fiador_dict | contrato_dic
            tipo = "(COM FIADOR)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato08-locatarios-fiador.docx")

        # Com dois Locatarios dois fiadores e caução
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 1 and self.varfiad.get() == 1 and self.var2fiador.get() == 1:
            my_dict = loca_dict | morador_dict | morador2_dict | fiador_dict | fiador_dict2 | contrato_dic
            tipo = "(COM FIADOR) (COM CAUÇÂO)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato09-locatarios-fiadores-caucao.docx")

        # Com dois Locatarios dois fiadores
        elif self.quantidadeLocatarios.get() == 2 and self.caucaoVar.get() == 0 and self.varfiad.get() ==1  and self.var2fiador.get() == 1:
            my_dict = loca_dict | morador_dict | morador2_dict | fiador_dict | fiador_dict2 | contrato_dic
            tipo = "(COM FIADOR)"
            doc = Document("F:/DJ/DP/PEDRO ALMEIDA/Programa_Contratos/Contratos/Programa de Contrato/docs/contrato10-locatarios-fiadores.docx")

        print(f"{self.quantidadeLocatarios.get()} -- {self.caucaoVar.get()} -- {self.varfiad.get()} -- {self.var2fiador.get()}")
        docx_replace(doc, **my_dict)

        new_name = self.nomeLoca.get().strip().split(" ")
        new_name = '_'.join(new_name)

        arquivo_name = f"CONTRATO_DE_LOCAÇÃO_{new_name}_X_{self.nomeMora1.get()}__{tipo}"

        doc.save(f"C:/Users/Usuario/Desktop/{arquivo_name}.docx")
        self.salve.configure(style='success.TButton')

    def insere_dados(self):
        self.nomeLoca.insert(0, "RUAN HENRIQUE DA SILVA FERREIRA")
        self.estcivLoca.insert(0, "SOLTEIRO")
        self.cpfLoca.insert(0, "502.699.448-26")
        self.rgLoca.insert(0, "43.432.324-0")
        self.org_ex_locador.insert(0, "SSP")
        self.estado_rg_locador.insert(0, "SP")
        self.cidad_locador.insert(0, "MIRASSOL")
        self.est_locador.insert(0, "SP")
        self.ruaLocador.insert(0, "RUA MATHEUS LEITE")
        self.bairroLocador.insert(0, "S. BERNARDO")
        self.cepLocador.insert(0, "15130-057")
        self.combLocador.insert(0, "Homem")
        self.nomeMora1.insert(0, "MARY LUCY")
        self.estcivMora.insert(0, "VIUVA")
        self.cpfMora.insert(0, "444.444.333-22")
        self.rgMora.insert(0, "22.333.444-5")
        self.org_ex_mora.insert(0, "SSP")
        self.estado_rg_mora.insert(0, "SP")
        self.cidad_mora.insert(0, "MIRASSOL")
        self.est_mora.insert(0, "SP")
        self.endImov.insert(0, "RUA RUY BARBOSA")
        self.numImov.insert(0, "4950")
        self.bairImov.insert(0, "CENTRO")
        self.municImov.insert(0, "MIRASSOL")
        self.estImov.insert(0, "SP")
        self.vencImov.insert(0, "05 (CINCO)")
        self.fimImov.entry.insert(0, "23/02/2024")
        self.prazo.insert(0, "12 meses")
        self.valorImov.insert(0, "R$ 1500,00")


        self.nomeMora2.insert(0, 'JULIO SEGUNDO')
        self.estcivMora2.insert(0, "CASADO")
        self.cpfMora2.insert(0, "777.555.999-55")
        self.rgMora2.insert(0, "55.444.666-4")
        self.org_ex_mora2.insert(0, "SSP")
        self.estado_rg_mora2.insert(0, "SP")
        self.cidad_mora2.insert(0, "MIRASSOL")
        self.est_mora2.insert(0, "SP")


            
        self.nomeFiador.insert(0, "JULIANA FREITAS")
        self.estcivFiador.insert(0, "VIUVA")
        self.cpfFiador.insert(0, "666.333.999-44")
        self.rgFiador.insert(0, "22.111.777-9")
        self.org_ex_fiador.insert(0, "SSP")
        self.estado_rg_fiador.insert(0, "SSP")
        self.cidad_fiador.insert(0, "MIRASSOL")
        self.ruaFiador.insert(0, "RUA SEM SAIDA")
        self.numeroFiador.insert(0, "3945")
        self.bairroFiador.insert(0, "SÂO FRANCISCO")
        self.est_fiador.insert(0, "SP")
        self.cepFiador.insert(0, "12345-069")

        self.nomeFiador2.insert(0, 'JULIO SEGUNDO')
        self.estcivFiador2.insert(0, "CASADO")
        self.cpfFiador2.insert(0, "777.555.999-55")
        self.rgFiador2.insert(0, "55.444.666-4")
        self.org_ex_fiador2.insert(0, "SSP")
        self.estado_rg_fiador2.insert(0, "SP")
